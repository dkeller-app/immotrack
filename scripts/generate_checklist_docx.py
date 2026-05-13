"""
Génère docs/audit/CHECKLIST-VALIDATION-V1.docx — version Word de la checklist.

Format :
  - Page de garde + sommaire
  - 12 sections avec headings hiérarchisés
  - Tableaux par sous-section : Action / Attendu / Vérif / KO / ☐ Statut / Notes
  - Cases à cocher Word natives (□) cliquables
  - Couleurs de section
  - En-tête + pied de page
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Données identiques à l'Excel (re-importées)
import sys, os
sys.path.insert(0, os.path.dirname(__file__))

# On reprend la même liste CHECKLIST que le script xlsx
CHECKLIST = [
    # (Section, Sous-section, N°, Action, Attendu, Où vérifier, KO si, durée)
    # PARTIE 1 — AUTO
    ("1. AUTO", "1.1 Tests Vitest", "1.1.1", "Ouvrir terminal + `npm run test:run`", "`Test Files 17 passed (17)` et `Tests 321 passed (321)`", "Sortie console", "Test rouge OU durée > 30s OU erreur module", 2),
    ("1. AUTO", "1.1 Tests Vitest", "1.1.2", "Vérifier absence de message rouge ✗", "Tous fichiers en vert ✓", "Couleur sortie console", "Un ✗ rouge → noter test/ligne/message", 1),
    ("1. AUTO", "1.1 Tests Vitest", "1.1.3", "Vérifier `Duration < 5s`", "Affiché en bas (~1-2s typique)", "Bas de sortie console", "Durée > 5000 ms = perf problème", 1),
    ("1. AUTO", "1.1 Tests Vitest", "1.1.4", "`npm test` (watch mode)", "Lance Vitest et attend modifications", "Console", "Erreur de boot ou crash", 1),
    ("1. AUTO", "1.2 HTTP Modules", "1.2.1", "Lancer http-server : `npx --yes http-server . -p 8766 -c-1 --silent &`", "Serveur démarré sans erreur", "Console terminal", "Port utilisé OU permission denied", 1),
    ("1. AUTO", "1.2 HTTP Modules", "1.2.2-16", "Vérifier 15 modules ES : css/main.css, js/main.js, js/core/*.js (10), js/components/*.js (3)", "Tous HTTP 200 avec tailles attendues (voir Excel)", "Sortie curl batch", "Un 404 → fichier manquant", 3),
    ("1. AUTO", "1.3 Console DevTools", "1.3.1", "Chrome → http://localhost:8766/index-test.html ; F12 ; Console", "Zéro erreur rouge", "DevTools Console", "Erreur JS bloquante", 2),
    ("1. AUTO", "1.3 Console DevTools", "1.3.2", "Message vert `[main.js] Sprint 3D chargé - 38 helpers...`", "Présent", "Console", "Absent → module ES n'a pas chargé", 1),
    ("1. AUTO", "1.3 Console DevTools", "1.3.3", "Network : `main.js` 200 type=module", "Vert", "DevTools Network", "Rouge → CSP/CORS/404", 1),
    ("1. AUTO", "1.3 Console DevTools", "1.3.4", "Taper `window.__IMMOTRACK_MODULE_BOOTSTRAP__`", "Objet retourné", "Console", "undefined", 1),
    ("1. AUTO", "1.3 Console DevTools", "1.3.5", "`window.__IMMOTRACK_MODULE_BOOTSTRAP__.helpersExposed.length`", "≥ 38", "Console", "< 38 → imports manquants", 1),
    ("1. AUTO", "1.3 Console DevTools", "1.3.6-8", "Vérifier `typeof window.escHtml/window._compute2044/window._auditEntry` === 'function'", "Tous `true`", "Console", "`false` → bug import", 1),

    # PARTIE 2 — SÉCU
    ("2. SÉCU XSS", "2.1 Helpers escHtml", "2.1.1-6", "Tester `escHtml('<x>')` → `\"&lt;x&gt;\"`, `escHtml('\"')` → `\"&quot;\"`, etc. (6 cas)", "Toutes valeurs exactes attendues", "Console", "Une valeur différente → escHtml cassé BLOQUANT", 3),
    ("2. SÉCU XSS", "2.1 Helpers escHtml", "2.1.7-8", "`_h\\`<b>${'<script>'}</b>\\`` → `<b>&lt;script&gt;</b>` + test `_raw`", "Comportement attendu", "Console", "Bug _h ou _raw", 2),
    ("2. SÉCU XSS", "2.2 XSS entité", "2.2.1-5", "Créer entité avec nom `<img src=x onerror=\"window.__xss=1\">` et tester partout", "`window.__xss` reste undefined ; aucune popup", "Console + Visuel", "`window.__xss === 1` → XSS active BLOQUANT V1", 5),
    ("2. SÉCU XSS", "2.3 XSS catégorie", "2.3.1-2", "Créer catégorie avec payload `<script>window.__xss2=1</script>` et utiliser", "`window.__xss2` reste undefined", "Console", "Exécution → BLOQUANT", 2),
    ("2. SÉCU XSS", "2.4 XSS locataire (PII)", "2.4.1-2", "Créer bail avec locataire `<img src=x onerror=alert(1)>` et naviguer", "Aucun alert/popup", "Visuel + Console", "Alert → BLOQUANT", 2),
    ("2. SÉCU XSS", "2.5 Validation saisie", "2.5.1-8", "Tester refs `<script>`, `F-001!@#$`, longues, normales + HC=99999, HC>CH", "Toasts erreur cohérents", "Toast", "Validation faussement permissive", 8),

    # PARTIE 3 — ARCHI
    ("3. ARCHI", "3.1 Fichiers", "3.1.1-3", "`ls -la css/main.css js/main.js js/core/ js/components/`", "Tous présents (10 core, 3 components)", "ls", "Fichier manquant", 1),
    ("3. ARCHI", "3.2 Tag git", "3.2.1", "`git tag | grep pre-modular-sprint2`", "Tag présent", "git", "Absent → rollback impossible", 1),
    ("3. ARCHI", "3.3 Pattern shadow", "3.3.1-3", "`window.escHtml === escHtml`, `window.openM === openM`, click onclick legacy", "true / true / modale s'ouvre", "Console + Visuel", "Idempotence cassée", 2),
    ("3. ARCHI", "3.4 Backward compat", "3.4.1-10", "Naviguer chaque onglet sidebar (Dashboard/Biens/Baux/Mouvements/Quittances/IRL/Régul/EDL/Paramètres/Export)", "Chaque page rend sans erreur", "Visuel", "Écran blanc ou erreur JS", 5),

    # PARTIE 4 — CONFORMITÉ
    ("4. CONFORMITÉ", "4.1 AUDIT-TRAIL", "4.1.1-9", "Saisir nom user → créer/modifier/supprimer mvt et bail → vérifier DB.auditTrail enrichi", "Entrées avec ts, userId, action, entityType, diff (sur update)", "Console", "auditTrail vide → hooks cassés", 5),
    ("4. CONFORMITÉ", "4.1 AUDIT-TRAIL", "4.1.10-11", "AUDIT COMPLÉTUDE v14.98 : créer immeuble + quittance + assurance + MRH + EDL → vérifier types couverts", "Tous types présents : entite, immeuble, logement, bail, quittance, assurance, mrh, edl, mouvement", "Console", "Type manquant → hook absent", 7),
    ("4. CONFORMITÉ", "4.2 LEGAL-2044", "4.2.1-9", "Carte 'Aide déclaration 2044' → sélectionner année + entité → Calculer + CSV", "Récap avec lignes 211/213/221-230/250, Résultat foncier, avertissement déficit si <0", "Visuel + CSV Excel", "Récap incomplet ou CSV mal formé", 5),
    ("4. CONFORMITÉ", "4.3 BILAN ANNUEL", "4.3.1-6", "Carte 'Bilan annuel' → entité + année → Générer", "Texte avec KPIs (revenus/charges/cash-flow/occupation) + détail par logement", "Visuel", "KPIs incohérents ou détail manquant", 4),
    ("4. CONFORMITÉ", "4.4 RGPD", "4.4.1-9", "Carte 'RGPD' → 3 boutons (rapport/export/plan effacement) + vérifier docs/legal/*.md existent", "Rapport texte, JSON portable art.20, plan d'effacement éligibilité 3 ans", "Visuel + Downloads", "Éligibilité incorrecte ou fichier non téléchargé", 8),
    ("4. CONFORMITÉ", "4.5 EXPORT-COMPTABLE FEC", "4.5.1-10", "Carte 'Compta' → 3 boutons FEC/Journal/Grand livre + vérification format DGFiP", "FEC 18 colonnes tab, date YYYYMMDD, montants virgule FR, GL000001 numérotation, débit=crédit équilibré", "Texte + Excel", "Format non conforme arrêté 29/07/2013", 10),
    ("4. CONFORMITÉ", "4.5 EXPORT-COMPTABLE FEC", "4.5.11", "TEST EXTERNE : envoyer FEC à expert-comptable", "Validation 'OK arrêté 29/07/2013 BOI-CF-IOR-60-40-20'", "Email retour", "Rejet par expert", 10),
    ("4. CONFORMITÉ", "4.6 IMPORT-CONCURRENTS", "4.6.1-4", "Tests console : `_mapRentila(...)`, `_mapBailFacile(...)`, `_mergeImport(DB, ...)` + dédup", "Objets structurés retournés, dédup fonctionnelle", "Console", "Throw ou doublons créés", 3),
    ("4. CONFORMITÉ", "4.7 MOBILE", "4.7.1-10", "DevTools Ctrl+Shift+M : tester 320 / 390 / 428 / 768 / 1280 px + interactions", "Layout responsive correct, touch ≥ 44px, modales 100dvh, sidebar burger <480px", "Visuel responsive", "Layout cassé ou bouton trop petit", 15),

    # PARTIE 5 — BUGS P1
    ("5. BUGS P1", "5.1 BUG-CHARGE-001", "5.1.1-5", "Créer mvts LEGAL-2044 → Régul → vérifier provisions ≠ 0", "Calcul correct legacy + LEGAL-2044", "Visuel Régul", "Provisions = 0 → régression", 5),
    ("5. BUGS P1", "5.2 BUG-DASH-001", "5.2.1-5", "Bail avec révision IRL → Dashboard avant/après → `_loyerHCAtDate(...)`", "Loyer historique correct à chaque date", "Visuel + Console", "Affiche mauvais HC", 5),
    ("5. BUGS P1", "5.3 DB-CORRUPT-FALLBACK", "5.3.1-3", "Corrompre localStorage + F5", "Toast erreur 12s + backup créé", "Visuel + Console", "App crash sans warning", 3),
    ("5. BUGS P1", "5.4 BUG-EQUIP-FILTER", "5.4.1-3", "Logement vacant → onglet Équipements", "Vacant inclus avec label", "Visuel select", "Vacant exclu → bug régression", 3),

    # PARTIE 6
    ("6. MONIT+CI", "6.1 Monitoring", "6.1.1-9", "Activer monitoring → installer capture → throw → vérifier capture + anonymisation + export/clear", "Erreur capturée avec hash userAgent (pas en clair)", "Console", "Capture cassée ou leak userAgent", 5),
    ("6. MONIT+CI", "6.2 CI GitHub", "6.2.1-4", "`.github/workflows/test.yml` + push test branche → workflow vert", "Workflow OK", "GitHub Actions", "Échec CI", 5),

    # PARTIE 7
    ("7. EMAIL-AUTO", "7. Email", "7.1-4", "Console : `_emailTypesSupportes()`, `_emailCompose(...)`, `DB.emailsSent` + UI buttons si présent", "10 types, objet structuré", "Console + UI", "Module absent ou throw", 5),

    # PARTIE 8
    ("8. DRIVE SYNC", "8. Drive", "8.1-8", "Connecter Drive → vérifier arborescence ImmoTrack/<entité>/<imm>/<log>/[9 sous-dossiers] + sync bidirectionnelle + tombstone", "Structure complète, push/pull OK, tombstones propagés", "Drive web + 2 devices", "Sync échoue ou résurrection", 20),

    # PARTIE 9
    ("9. PARCOURS", "9.1 Bailleur particulier", "9.1.1-14", "Parcours complet : entité → immeuble → logement → bail → 12 mvts → quittances → régul → 2044 → bilan → FEC", "Tous fonctionnels + audit-trail ≥ 16", "Visuel + Console", "Étape cassée ou audit manquant", 25),
    ("9. PARCOURS", "9.2 IRL complet", "9.2.1-7", "DPE → calcul révision → lettre → valider envoi → valider IRL → vérif Dashboard temporel + gel DPE F", "Cycle IRL complet + loi Climat 2021", "Visuel + Console", "Cycle interrompu ou DPE F ignoré", 10),
    ("9. PARCOURS", "9.3 EDL complet", "9.3.1-7", "EDL entrée → 3 pièces + photos + signatures → save → audit hook → PDF → Drive", "EDL complet + audit + PDF + Drive sync", "Visuel + Downloads + Drive", "Étape cassée", 10),

    # PARTIE 10
    ("10. ERRORS", "10. Console clean", "10.1.1-3", "F5 + Console errors only + Network failed", "Tout vide", "DevTools", "Erreurs présentes", 2),

    # PARTIE 11
    ("11. AVANCÉ", "11.1 Audit complétude", "11.1.1", "Snippet console : check tous types audit-trail couverts", "Tous ✅", "Console", "Type ⚠️", 2),
    ("11. AVANCÉ", "11.2 Performance", "11.2.1-3", "DevTools Performance + Memory + localStorage size", "FPS > 30, heap < 50 MB, localStorage < 5 MB", "DevTools", "Bottleneck ou fuite mémoire", 5),
    ("11. AVANCÉ", "11.3 Accessibilité", "11.3.1", "Lighthouse Accessibility audit", "Score > 80", "Lighthouse", "Score < 60 → améliorer", 3),

    # PARTIE 12
    ("12. BASCULE", "12.1 Backups obligatoires", "12.1.1-4", "Backup index.html + export JSON DB + tag git pre-bascule + vérif tag pre-modular-sprint2", "Tous backups effectués", "ls + git tag", "Backup oublié", 3),
    ("12. BASCULE", "12.2 Bascule commits", "12.2.1-2", "Liste 19 commits puis reproduction granulaire sur index.html", "1 commit prod / sprint sandbox", "git log + prod", "Commit géant", 60),
]

# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

doc = Document()

# Marges
sections_doc = doc.sections
for section in sections_doc:
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

# ━━━━ Page de garde ━━━━
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("CHECKLIST DE VALIDATION V1 COMMERCIALE\n\nImmoTrack v14.98")
run.font.size = Pt(24)
run.font.bold = True
run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x78)

doc.add_paragraph()
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub_run = sub.add_run("Document généré le 2026-05-12 — Audit post-marathon 4 sprints")
sub_run.font.size = Pt(11)
sub_run.font.italic = True
sub_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph()
intro = doc.add_paragraph()
intro_run = intro.add_run(
    "Cette checklist couvre 213 points de contrôle sur 12 sections pour valider la sandbox v14.98 "
    "avant bascule en production. Durée totale estimée : 3-4 heures pour validation exhaustive.\n\n"
    "Pour chaque point, cocher ☑ dans la colonne Statut quand validé. "
    "Noter en colonne « Notes » tout écart entre résultat attendu et constaté.\n\n"
    "Préparation : "
)
intro_run.font.size = Pt(11)
cmd = intro.add_run("cd C:/Users/Did_K/Desktop/Immo && npx --yes http-server . -p 8766 -c-1 --silent &")
cmd.font.name = "Consolas"
cmd.font.size = Pt(10)
intro.add_run(
    "\n\nOuvrir Chrome → http://localhost:8766/index-test.html → F12 onglet Console.\n\n"
    "En cas d'incident critique : "
).font.size = Pt(11)
cmd2 = intro.add_run("git reset --hard pre-modular-sprint2")
cmd2.font.name = "Consolas"
cmd2.font.size = Pt(10)

doc.add_page_break()

# ━━━━ Sommaire ━━━━
toc_title = doc.add_paragraph()
toc_run = toc_title.add_run("SOMMAIRE")
toc_run.font.size = Pt(20)
toc_run.font.bold = True
toc_run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x78)

sections_list = [
    "1. AUTO — Tests automatiques (Vitest, HTTP, console)",
    "2. SÉCU XSS — Helpers + tests injection HTML",
    "3. ARCHI — Modules ES + backward compat",
    "4. CONFORMITÉ — Audit-trail, 2044, RGPD, Bilan, FEC, Mobile",
    "5. BUGS P1 — 4 bugs fixés à reproduire",
    "6. MONIT+CI — Monitoring opt-in + GitHub Actions",
    "7. EMAIL-AUTO — Module commits parallèles",
    "8. DRIVE SYNC — Bidirectionnel + arborescence",
    "9. PARCOURS — 3 scénarios intégration complets",
    "10. ERRORS — Console clean",
    "11. AVANCÉ — Perf + a11y + audit transversal",
    "12. BASCULE — Backups + 19 commits ordonnés",
]

for s in sections_list:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.add_run("☐  ").font.size = Pt(11)
    r = p.add_run(s)
    r.font.size = Pt(11)

doc.add_page_break()

# ━━━━ Helper : créer un tableau de checklist par sous-section ━━━━

def section_header(text, color_rgb=(0x1F, 0x4E, 0x78)):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = RGBColor(*color_rgb)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    # Bordure inférieure via XML
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '8')
    bottom.set(qn('w:color'), '1F4E78')
    pBdr.append(bottom)
    pPr.append(pBdr)

def subsection_header(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)

def add_checklist_table(rows):
    """rows = liste de tuples (num, action, attendu, ou_verif, ko, duree)"""
    table = doc.add_table(rows=len(rows) + 1, cols=6)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Headers
    hdr = table.rows[0].cells
    headers = ["N°", "Action", "Résultat attendu", "Où vérifier / KO si", "☐ OK", "Notes"]
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        r = p.add_run(h)
        r.font.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        # Fond bleu via XML
        tcPr = hdr[i]._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '1F4E78')
        tcPr.append(shd)

    # Largeurs colonnes (en cm)
    widths_cm = [1.2, 4.5, 4.5, 4.0, 1.0, 3.0]
    for col_idx, w in enumerate(widths_cm):
        for cell in table.columns[col_idx].cells:
            cell.width = Cm(w)

    # Data rows
    for row_idx, (num, action, attendu, ou_verif, ko, duree) in enumerate(rows, start=1):
        cells = table.rows[row_idx].cells

        # N°
        p = cells[0].paragraphs[0]
        r = p.add_run(num)
        r.font.size = Pt(9)
        r.font.bold = True

        # Action
        p = cells[1].paragraphs[0]
        r = p.add_run(action)
        r.font.size = Pt(9)
        # Si action contient code, mettre en monospace pour la partie code
        # (ici simplifié, on garde tout en sans-serif pour lisibilité)

        # Résultat attendu
        p = cells[2].paragraphs[0]
        r = p.add_run(attendu)
        r.font.size = Pt(9)

        # Où vérifier + KO
        p = cells[3].paragraphs[0]
        r = p.add_run(f"📍 {ou_verif}")
        r.font.size = Pt(8)
        r.font.color.rgb = RGBColor(0x33, 0x66, 0x99)
        p.add_run("\n")
        r2 = p.add_run(f"❌ {ko}")
        r2.font.size = Pt(8)
        r2.font.color.rgb = RGBColor(0xCC, 0x33, 0x33)

        # ☐ checkbox (caractère Unicode case à cocher)
        p = cells[4].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run("☐")
        r.font.size = Pt(16)

        # Notes (vide)
        cells[5].text = ""

        # Vertical alignment
        for c in cells:
            c.vertical_alignment = WD_ALIGN_VERTICAL.TOP

# ━━━━ Génération du contenu par section ━━━━

current_section = None
current_subsection = None
buffer_rows = []  # rows pour la subsection en cours

def flush_buffer():
    global buffer_rows
    if buffer_rows:
        add_checklist_table(buffer_rows)
        buffer_rows = []

for item in CHECKLIST:
    section, subsection, num, action, attendu, ou_verif, ko, duree = item

    # Changement de section → nouveau heading + flush
    if section != current_section:
        flush_buffer()
        if current_section is not None:
            doc.add_page_break()
        section_header(section)
        current_section = section
        current_subsection = None

    # Changement de sous-section → flush + nouveau heading
    if subsection != current_subsection:
        flush_buffer()
        subsection_header(subsection)
        current_subsection = subsection

    buffer_rows.append((num, action, attendu, ou_verif, ko, duree))

flush_buffer()

# ━━━━ Page finale : Récap ━━━━
doc.add_page_break()

p = doc.add_paragraph()
r = p.add_run("RÉCAP FINAL")
r.font.size = Pt(20)
r.font.bold = True
r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x78)

recap_items = [
    "☐ Partie 1 (auto) : 321 tests Vitest passent, 15 modules HTTP 200, console clean",
    "☐ Partie 2 (SÉCU) : 0 XSS exploitable, validation ref OK, helpers escHtml conformes",
    "☐ Partie 3 (Archi) : Modules ES chargés, pattern shadow idempotent, navigation backward OK",
    "☐ Partie 4 (Conformité) : AUDIT-TRAIL, LEGAL-2044, RGPD, Bilan, FEC, Mobile — tous fonctionnels",
    "☐ Partie 5 (Bugs P1) : 4 bugs reproduits et vérifiés fixés",
    "☐ Partie 6 (Monitoring + CI) : Logs anonymes, CI verte sur GitHub Actions",
    "☐ Partie 7 (EMAIL-AUTO) : Tests Vitest OK, UI à explorer si présente",
    "☐ Partie 8 (Drive) : Sync bidirectionnelle, arborescence 9 sous-dossiers, tombstones",
    "☐ Partie 9 (Parcours) : Bailleur particulier + IRL + EDL complets sans accroc",
    "☐ Partie 10 (Errors/warnings) : Console propre, aucune erreur rouge",
    "☐ Partie 11 (Audits avancés) : Performance + a11y dans cibles, audit complétude OK",
    "☐ Partie 12 (Bascule prod) : Backups effectués, méthode commits granulaires validée",
]

for item in recap_items:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(item)
    r.font.size = Pt(11)

# Conclusion
doc.add_paragraph()
final = doc.add_paragraph()
final.paragraph_format.space_before = Pt(12)
r1 = final.add_run("Si tous ✅ → ")
r1.font.size = Pt(11)
r2 = final.add_run("V1 PRÊTE POUR BÊTA PRIVÉE")
r2.font.size = Pt(11)
r2.font.bold = True
r2.font.color.rgb = RGBColor(0x00, 0x88, 0x00)
final.add_run(".").font.size = Pt(11)

final2 = doc.add_paragraph()
r3 = final2.add_run("Si ≥ 1 ❌ critique → ")
r3.font.size = Pt(11)
r4 = final2.add_run("RAPPORT DÉTAILLÉ")
r4.font.size = Pt(11)
r4.font.bold = True
r4.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
final2.add_run(" (ligne, étape, attendu vs constaté) → fix avant bascule.").font.size = Pt(11)

# Sauvegarde
import os
out_path = "C:/Users/Did_K/Desktop/Immo/docs/audit/CHECKLIST-VALIDATION-V1.docx"
doc.save(out_path)

print(f"OK Fichier genere : {out_path}")
print(f"   Sections : 12")
print(f"   Sous-sections : {len(set((s, ss) for s, ss, *_ in CHECKLIST))}")
print(f"   Lignes data : {len(CHECKLIST)}")
print(f"   Taille : {os.path.getsize(out_path)} octets")
